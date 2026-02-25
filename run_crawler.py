"""
Main crawler entry point.

Usage:
    python run_crawler.py [--storage PATH] [--output PATH] [--generate-samples]
    python run_crawler.py search "query" [--limit N]
"""

import argparse
import random
import string
import sys
from datetime import datetime
from pathlib import Path
from typing import Any

from structlog import get_logger

from app.application.use_cases import CrawlUseCase, SearchUseCase
from app.config import get_config
from app.infrastructure.logger import setup_logging


def parse_args():
    """Parse command line arguments."""
    parser = argparse.ArgumentParser(
        description="Document crawler with full-text indexing",
        formatter_class=argparse.RawDescriptionHelpFormatter,
    )

    subparsers = parser.add_subparsers(dest="command", help="Commands")

    crawl_parser = subparsers.add_parser(
        "crawl",
        help="Crawl storage directory and index documents",
    )
    crawl_parser.add_argument(
        "--storage",
        type=Path,
        help="Path to storage directory with files to crawl",
    )
    crawl_parser.add_argument(
        "--output",
        type=Path,
        help="Path to output CSV file",
    )
    crawl_parser.add_argument(
        "--limit",
        type=int,
        help="Maximum number of files to process",
    )
    crawl_parser.add_argument(
        "--no-archives",
        action="store_true",
        help="Don't extract archives",
    )

    search_parser = subparsers.add_parser(
        "search",
        help="Search indexed documents",
    )
    search_parser.add_argument(
        "query",
        type=str,
        help="Search query (FTS5 syntax)",
    )
    search_parser.add_argument(
        "--limit",
        type=int,
        default=20,
        help="Maximum results to return",
    )
    search_parser.add_argument(
        "--type",
        dest="doc_type",
        choices=["pdf", "docx", "xlsx", "txt"],
        help="Filter by document type",
    )

    subparsers.add_parser(
        "stats",
        help="Show database statistics",
    )

    samples_parser = subparsers.add_parser(
        "generate-samples",
        help="Generate sample files for testing",
    )
    samples_parser.add_argument(
        "--output",
        type=Path,
        default=Path("tests/fixtures/samples"),
        help="Output directory for samples",
    )
    samples_parser.add_argument(
        "--count",
        type=int,
        default=5,
        help="Number of sample files to generate per type",
    )

    return parser.parse_args()


def generate_sample_files(output_dir: Path, count: int) -> None:
    """Generate sample files of various formats for testing."""

    output_dir.mkdir(parents=True, exist_ok=True)
    logger = get_logger("run_crawler.py")

    logger.info(f"Generating {count} sample files in {output_dir}")

    for i in range(count):
        content = "\n".join([
            f"Sample text file {i + 1}",
            f"Generated at: {datetime.now()}",
            "".join(random.choices(string.ascii_letters, k=100)),
        ])
        (output_dir / f"sample_{i + 1}.txt").write_text(content)

    for i in range(count):
        content = f"""# Sample Markdown {i + 1}

Generated at: {datetime.now()}

## Section 1
Lorem ipsum dolor sit amet.

## Section 2
Consectetur adipiscing elit.

- List item 1
- List item {i + 1}
"""
        (output_dir / f"sample_{i + 1}.md").write_text(content)

    for i in range(count):
        import json

        data: dict[str, Any] = {
            "id": i + 1,
            "name": f"Sample {i + 1}",
            "timestamp": datetime.now().isoformat(),
            "values": random.sample(range(1, 100), 5),
        }

        (output_dir / f"sample_{i + 1}.json").write_text(
            json.dumps(data, indent=2)
        )

    try:
        from reportlab.pdfgen import canvas

        for i in range(count):
            pdf_path = output_dir / f"sample_{i + 1}.pdf"
            c = canvas.Canvas(str(pdf_path))
            c.drawString(100, 750, f"Sample PDF {i + 1}")
            c.drawString(100, 735, f"Generated at: {datetime.now()}")
            c.save()
    except ImportError:
        logger.info("reportlab not installed, skipping PDF generation")

    try:
        import xlwt  # pyright: ignore[reportMissingTypeStubs]

        logger.info("Generating XLS samples...")
        for i in range(count):
            workbook = xlwt.Workbook(encoding="utf-8")
            sheet = workbook.add_sheet("Sheet1")  # type: ignore
            sheet.write(0, 0, f"Sample XLS {i + 1}")  # type: ignore
            sheet.write(  # type: ignore
                1,
                0,
                f"Generated at: {datetime.now()}",
            )
            for j in range(5):
                sheet.write(j + 2, 0, f"Value {j}")  # type: ignore
                sheet.write(j + 2, 1, j * 100)  # type: ignore

            output_path = output_dir / f"sample_{i + 1}.xls"
            workbook.save(str(output_path))  # type: ignore
            logger.debug(f"Generated {output_path}")

    except ImportError:
        logger.warning("xlwt not installed, skipping XLS generation")

    try:
        from docx import Document

        for i in range(count):
            doc = Document()
            doc.add_heading(f"Sample DOCX {i + 1}", 0)
            doc.add_paragraph(f"Generated at: {datetime.now()}")
            doc.add_paragraph("Lorem ipsum dolor sit amet.")
            doc.save(str(output_dir / f"sample_{i + 1}.docx"))
    except ImportError:
        logger.info("python-docx not installed, skipping DOCX generation")

    try:
        import pandas as pd

        for i in range(count):
            df = pd.DataFrame({
                "ID": range(1, 6),
                "Value": random.sample(range(1, 100), 5),
                "Text": ["A", "B", "C", "D", "E"],
            })
            df.to_excel(  # pyright: ignore[reportUnknownMemberType]
                output_dir / f"sample_{i + 1}.xlsx", index=False
            )
    except ImportError:
        logger.info("pandas not installed, skipping XLSX generation")

    logger.info(f"Generated sample files in {output_dir}")


def main() -> int:
    """Main entry point."""
    args = parse_args()

    config = get_config()
    setup_logging(config.logger_adapter)
    logger = get_logger("run_crawler.py")

    if args.command == "generate-samples":
        generate_sample_files(args.output, args.count)
        return 0

    if args.command == "crawl":
        logger.info("=== START CRAWL ===")

        if args.storage:
            config.crawler.storage_path = args.storage
        if args.output:
            config.crawler.output_csv_path = args.output
        if args.no_archives:
            config.crawler.extract_archives = False

        try:
            crawl_use_case = CrawlUseCase(config)
            stats = crawl_use_case.execute(limit=args.limit)

            print("\n" + "=" * 60)
            print("CRAWL COMPLETED".center(60))
            print("=" * 60)
            print(f"Files found:     {stats['files_found']}")
            print(f"Files processed: {stats['files_processed']}")
            print(f"Files failed:    {stats['files_failed']}")
            print(f"Files skipped:   {stats['files_skipped']}")
            print(f"Archives:        {stats['archives_extracted']}")
            print(f"Duration:        {stats['duration_seconds']:.2f} seconds")
            print("-" * 60)
            print("Database stats:")
            db_stats = stats.get("database", {})
            print(f"  Total docs:     {db_stats.get('total_documents', 0)}")
            print(
                f"  With text:      {db_stats.get('documents_with_text', 0)}"
            )
            print(
                f"  Total size:     {db_stats.get('total_size_mb', 0):.2f} MB"
            )
            print("=" * 60)

            return 0

        except Exception as e:
            logger.exception("Crawl failed")
            print(f"Error: {e}", file=sys.stderr)
            return 1

    elif args.command == "search":
        logger.info(f"=== SEARCH: {args.query} ===")

        try:
            search_use_case = SearchUseCase(config)
            results = search_use_case.search(
                query=args.query,
                limit=args.limit,
            )

            print("\n" + "=" * 60)
            print(f"SEARCH RESULTS: '{args.query}'".center(60))
            print("=" * 60)
            print(f"Found {results['total']} documents")
            print("-" * 60)

            for i, doc in enumerate(results["results"], 1):
                print(f"{i}. {doc['file_name']}")
                print(f"   Path: {doc['path']}")
                print(f"   Type: {doc['doc_type']}")
                if doc["from_archive"]:
                    print(f"   Archive: {doc['archive_path']}")
                if doc["text_preview"]:
                    print(f"   Preview: {doc['text_preview']}")
                print()

            print("=" * 60)
            return 0

        except Exception as e:
            logger.exception("Search failed")
            print(f"Error: {e}", file=sys.stderr)
            return 1

    elif args.command == "stats":
        try:
            search_use_case = SearchUseCase(config)
            stats = search_use_case.get_stats()

            print("\n" + "=" * 60)
            print("DATABASE STATISTICS".center(60))
            print("=" * 60)
            print(f"Total documents:     {stats['total_documents']}")
            print(f"With extracted text: {stats['documents_with_text']}")
            print(f"Without text:        {stats['documents_without_text']}")
            print(f"Total size:          {stats['total_size_mb']:.2f} MB")
            print("-" * 60)
            print("By document type:")
            for doc_type, count in stats["by_type"].items():
                print(f"  {doc_type}: {count}")
            print("=" * 60)

            return 0

        except Exception as e:
            logger.exception("Failed to get stats")
            print(f"Error: {e}", file=sys.stderr)
            return 1

    else:
        print("Please specify a command. Use --help for usage.")
        return 1


if __name__ == "__main__":
    sys.exit(main())
