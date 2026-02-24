"""
DOCX document parser using python-docx.
"""

from pathlib import Path
from typing import ClassVar

from docx import Document as DocxDocument
from docx.opc.exceptions import PackageNotFoundError
from structlog import get_logger

from app.domain.entities import DocumentType

from .base_parser import BaseParser
from .interfaces import ParseError


class DOCXParser(BaseParser):
    """Parser for DOCX documents."""

    SUPPORTED_EXTENSIONS: ClassVar[set[str]] = {"docx"}
    DOCUMENT_TYPE: ClassVar[DocumentType] = DocumentType.DOCX

    def _extract_text(self, file_path: Path) -> str:
        """
        Extract text from DOCX using python-docx.

        Args:
            file_path: Path to DOCX file.

        Returns:
            Extracted text from paragraphs and tables.

        Raises:
            ParseError: If DOCX cannot be read.
        """
        try:
            doc = DocxDocument(str(file_path))

            text_parts: list[str] = []

            for para in doc.paragraphs:
                if para.text:
                    text_parts.append(para.text)

            for table in doc.tables:
                for row in table.rows:
                    row_text: list[str] = []
                    for cell in row.cells:
                        if cell.text:
                            row_text.append(cell.text)
                    if row_text:
                        text_parts.append(" | ".join(row_text))

            return "\n".join(text_parts)

        except PackageNotFoundError as e:
            raise ParseError(
                "Invalid DOCX package (corrupted or not a DOCX file)",
                file_path=file_path,
                original_error=e,
            ) from e
        except FileNotFoundError as e:
            raise ParseError("File not found", file_path=file_path) from e
        except PermissionError as e:
            raise ParseError("Permission denied", file_path=file_path) from e
        except Exception as e:
            raise ParseError(
                f"DOCX parsing failed: {e}",
                file_path=file_path,
                original_error=e,
            ) from e

    def extract_metadata(self, file_path: Path) -> dict[str, str | int | None]:
        """
        Extract DOCX metadata.

        Returns:
            Dict with author, creation date, etc.
        """
        try:
            doc = DocxDocument(str(file_path))

            metadata: dict[str, str | int | None] = {}

            if doc.core_properties:
                props = doc.core_properties
                if props.author:
                    metadata["author"] = props.author
                if props.created:
                    metadata["created"] = props.created.isoformat()
                if props.modified:
                    metadata["modified"] = props.modified.isoformat()
                if props.title:
                    metadata["title"] = props.title
                if props.subject:
                    metadata["subject"] = props.subject

            return metadata

        except Exception as e:
            logger = get_logger(__name__)
            logger.warning(
                "Failed to extract DOCX metadata",
                file=str(file_path),
                error=str(e),
            )
            return {}
